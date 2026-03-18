
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import io

class DOCXGenerator:
    @staticmethod
    def generate(data, style_config=None):
        document = Document()
        
        # Style Configuration
        scale = 1.0
        if style_config and 'font_scale' in style_config:
            try:
                scale = float(style_config['font_scale'])
            except:
                scale = 1.0
        
        # Helper to scale font size
        def s(pt_size):
            return max(6, pt_size * scale) # Minimum 6pt safety

        # 1. Page Setup (A4)
        section = document.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

        # Helper to add styled paragraph
        def add_para(text, style='Normal', bold=False, size=11, color=None, alignment=None, space_after=Pt(6)):
            p = document.add_paragraph()
            # Do NOT scale spacing for now to prevent collapse issues
            p.paragraph_format.space_after = space_after
            if alignment:
                p.alignment = alignment
            
            run = p.add_run(text)
            run.font.name = 'Calibri'
            # Scale the font size
            run.font.size = Pt(s(size))
            run.bold = bold
            if color:
                run.font.color.rgb = color
            return p

        # 2. Header (Name & Contact)
        p_info = data.get('personal_info')
        if not p_info:
            # Fallback if personal_info is completely missing
            add_para("Your Name", size=24, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            full_name = getattr(p_info, 'full_name', '') or "Your Name"
            
            # Name
            add_para(full_name, size=24, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
            
            # Contact Line
            contact_parts = []
            email = getattr(p_info, 'email', '')
            if email: contact_parts.append(email)
            
            phone = getattr(p_info, 'phone', '')
            if phone: contact_parts.append(phone)
            
            linkedin = getattr(p_info, 'linkedin', '')
            if linkedin: contact_parts.append(linkedin)
            
            location = getattr(p_info, 'location', '')
            if location: contact_parts.append(location)
            
            if contact_parts:
                add_para(" | ".join(contact_parts), size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(100, 100, 100), space_after=Pt(20))

            # 3. Summary
            summary = getattr(p_info, 'summary', '')
            if summary:
                add_para("SUMMARY", size=12, bold=True, color=RGBColor(59, 130, 246)) # Blue accent
                # Basic content:
                # Strip HTML tags if present (since summary might be HTML from AI)
                try:
                    import re
                    clean_summary = re.sub('<[^<]+?>', '', summary)
                except:
                    clean_summary = summary
                
                add_para(clean_summary, size=11, space_after=Pt(18))

        # 4. Experience
        experiences = data.get('experiences', [])
        if experiences:
            add_para("EXPERIENCE", size=12, bold=True, color=RGBColor(59, 130, 246))
            for exp in experiences:
                if not exp: continue
                # Title & Date
                title_line = document.add_paragraph()
                title_line.paragraph_format.space_after = Pt(2) # Fixed small spacing
                
                job_title = getattr(exp, 'job_title', '') or "Job Title"
                run_title = title_line.add_run(job_title)
                run_title.font.name = 'Calibri'
                run_title.font.size = Pt(s(11))
                run_title.bold = True
                
                company = getattr(exp, 'company', '')
                company_text = f" at {company}" if company else ""
                
                date_text = ""
                start_date = getattr(exp, 'start_date', '')
                if start_date:
                    date_text = f" | {start_date}"
                    end_date = getattr(exp, 'end_date', '')
                    if end_date:
                        date_text += f" - {end_date}"
                    else:
                        date_text += " - Present"
                
                run_company = title_line.add_run(company_text + date_text)
                run_company.font.name = 'Calibri'
                run_company.font.size = Pt(s(11))
                run_company.italic = True
                
                description = getattr(exp, 'description', '')
                if description:
                    add_para(description, size=11, space_after=Pt(12))

        # 5. Education
        education = data.get('education', [])
        if education:
            add_para("EDUCATION", size=12, bold=True, color=RGBColor(59, 130, 246))
            for edu in education:
                degree = getattr(edu, 'degree', '') or 'Degree'
                institution = getattr(edu, 'institution', '') or 'Institution'
                year = getattr(edu, 'year', '')
                
                text = f"{degree} - {institution}"
                if year:
                    text += f" ({year})"
                add_para(text, size=11, space_after=Pt(6))

        # 6. Skills
        skills = data.get('skills', [])
        if skills:
            add_para("SKILLS", size=12, bold=True, color=RGBColor(59, 130, 246))
            final_skills = []
            for s_item in skills: # Renamed loop var to avoid conflict with s()
                if isinstance(s_item, dict):
                    final_skills.append(s_item.get('name', ''))
                elif hasattr(s_item, 'name'):
                    final_skills.append(s_item.name)
                else:
                    final_skills.append(str(s_item))
                    
            add_para(", ".join(filter(None, final_skills)), size=11, space_after=Pt(18))

        # Save to buffer
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer
