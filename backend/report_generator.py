"""
Report Generator Module for ResuMate AI Resume Analyzer
Generates professional, branded PDF and Word (.docx) career intelligence reports.
"""

import io
import json
from datetime import datetime


# ═══════════════════════════════════════════════════════════════════════
#  PDF REPORT GENERATION (xhtml2pdf)
# ═══════════════════════════════════════════════════════════════════════

def generate_pdf_report(analysis_data, filename=None):
    """Generate a premium branded PDF career intelligence report."""
    from xhtml2pdf import pisa

    if isinstance(analysis_data, str):
        analysis_data = json.loads(analysis_data)

    a = analysis_data
    now = datetime.now().strftime("%B %d, %Y")
    time_str = datetime.now().strftime("%I:%M %p")
    file_label = filename or "Resume"
    score = a.get("opportunity_score", 0) or 0
    score_color = "#059669" if score >= 80 else "#e67e22" if score >= 60 else "#e74c3c"
    score_label = "Excellent" if score >= 80 else "Good" if score >= 60 else "Needs Work"

    sal_data = a.get("salary_estimation_inr", {}) or {}
    sal_min_raw = sal_data.get("min")
    sal_max_raw = sal_data.get("max")

    def _esc(text):
        if not text:
            return ""
        return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def _format_salary(val):
        if not val:
            return "--"
        num = int(float(val))
        if num >= 100000:
            return f"Rs.{num / 100000:.1f}L"
        if num >= 1000:
            return f"Rs.{num // 1000}K"
        return f"Rs.{num}"

    def _pills(items, fg="#4f46e5", bg="#eef2ff"):
        if not items:
            return '<p class="empty">None detected</p>'
        return " ".join(
            f'<span style="display:inline-block;background:{bg};color:{fg};'
            f'padding:3px 10px;border-radius:10px;font-size:10px;font-weight:700;'
            f'margin:2px 2px;">{_esc(s)}</span>' for s in items
        )

    def _bullets(items, icon="&#8226;", color="#334155"):
        if not items:
            return '<p class="empty">None detected</p>'
        return "".join(
            f'<p style="margin:4px 0 4px 14px;font-size:11px;color:{color};line-height:1.6;">'
            f'<span style="color:{color};margin-right:6px;">{icon}</span>{_esc(it)}</p>'
            for it in items
        )

    def _numbered(items, accent="#4f46e5"):
        if not items:
            return '<p class="empty">No plan generated</p>'
        return "".join(
            f'<p style="margin:5px 0 5px 14px;font-size:11px;line-height:1.6;">'
            f'<span style="background:{accent};color:white;padding:1px 6px;border-radius:8px;'
            f'font-size:9px;font-weight:800;margin-right:8px;">{i+1}</span>{_esc(step)}</p>'
            for i, step in enumerate(items)
        )

    # ── JOB ROLES TABLE ──
    roles_html = ""
    for jr in (a.get("job_roles") or []):
        pct = jr.get("match_percentage", 0)
        c = "#059669" if pct >= 80 else "#e67e22" if pct >= 60 else "#e74c3c"
        roles_html += (
            f'<tr><td style="padding:7px 10px;border-bottom:1px solid #f1f5f9;font-size:11px;font-weight:700;">{_esc(jr.get("role",""))}</td>'
            f'<td style="padding:7px 10px;border-bottom:1px solid #f1f5f9;font-size:11px;font-weight:800;color:{c};text-align:center;">{pct}%</td>'
            f'<td style="padding:7px 10px;border-bottom:1px solid #f1f5f9;font-size:10px;color:#64748b;">{_esc(jr.get("reason",""))}</td></tr>'
        )

    # ── INDUSTRY TABLE ──
    ind_html = ""
    for ind in (a.get("industry_fit") or []):
        ind_html += (
            f'<tr><td style="padding:6px 10px;border-bottom:1px solid #f1f5f9;font-size:11px;font-weight:700;">{_esc(ind.get("industry",""))}</td>'
            f'<td style="padding:6px 10px;border-bottom:1px solid #f1f5f9;font-size:11px;font-weight:800;color:#0891b2;text-align:center;">{ind.get("fit_percentage",0)}%</td></tr>'
        )

    # ── CERTIFICATIONS ──
    certs_html = ""
    for cert in (a.get("certification_recommendations") or []):
        if isinstance(cert, str):
            certs_html += f'<p style="margin:3px 0 3px 14px;font-size:11px;">&#8226; {_esc(cert)}</p>'
        else:
            n = _esc(cert.get("name", ""))
            p_str = f' ({_esc(cert.get("provider", ""))})' if cert.get("provider") else ""
            certs_html += f'<p style="margin:3px 0 3px 14px;font-size:11px;">&#8226; <b>{n}</b>{p_str}</p>'

    html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"/>
<style>
@page {{
    size: A4;
    margin: 1.2cm 1cm 2cm 1.5cm;
    @frame footer_frame {{
        -pdf-frame-content: page_footer;
        bottom: 0; left: 0; right: 0; height: 1.8cm;
    }}
}}
body {{ font-family: Helvetica, Arial, sans-serif; color: #1e293b; font-size: 11px; line-height: 1.6; margin: 0; padding: 0; }}

/* HEADER */
.header {{
    background-color: #312e81;
    padding: 28px 24px 22px 24px;
    color: white;
}}
.header-brand {{ font-size: 11px; font-weight: 800; letter-spacing: 3px; text-transform: uppercase; color: #a5b4fc; margin: 0 0 2px 0; }}
.header-title {{ font-size: 24px; font-weight: 800; margin: 0 0 4px 0; letter-spacing: -0.3px; }}
.header-sub {{ font-size: 10px; color: #c7d2fe; margin: 0; }}

/* BODY */
.body-wrap {{ padding: 20px 18px 20px 22px; }}

/* META BAR */
.meta {{ background: #f8fafc; border: 1px solid #e2e8f0; border-radius: 6px; padding: 10px 16px; margin-bottom: 20px; }}
.meta span {{ font-size: 10px; color: #64748b; margin-right: 18px; }}
.meta b {{ color: #334155; }}

/* SECTIONS */
.sec {{ margin-bottom: 16px; page-break-inside: avoid; }}
.sec-head {{
    font-size: 13px; font-weight: 800; color: #312e81;
    border-left: 3px solid #6366f1; padding-left: 10px;
    margin-bottom: 8px;
}}
.empty {{ color: #94a3b8; font-style: italic; font-size: 10px; margin: 4px 0 4px 14px; }}

/* SCORE BOX */
.score-wrap {{ text-align: center; padding: 12px 0; }}
.score-box {{
    display: inline-block; background: {score_color}; color: white;
    font-size: 32px; font-weight: 900; padding: 14px 28px; border-radius: 12px;
    letter-spacing: -1px;
}}
.score-label {{ font-size: 10px; color: {score_color}; font-weight: 700; margin-top: 6px; }}

/* TABLES */
table {{ width: 100%; border-collapse: collapse; }}
th {{
    background: #f1f5f9; color: #475569; font-size: 9px; font-weight: 800;
    text-transform: uppercase; letter-spacing: 0.5px;
    padding: 7px 10px; text-align: left; border-bottom: 2px solid #cbd5e1;
}}

/* SALARY */
.sal-box {{ text-align: center; padding: 4px 0; }}
.sal-min {{ font-size: 18px; font-weight: 900; color: #d97706; }}
.sal-max {{ font-size: 18px; font-weight: 900; color: #059669; }}
.sal-arrow {{ font-size: 16px; color: #cbd5e1; margin: 0 10px; }}
</style>
</head>
<body>

<!-- HEADER -->
<div class="header">
    <p class="header-brand">ResuMate AI</p>
    <p class="header-title">Career Intelligence Report</p>
    <p class="header-sub">Deep AI-Powered Career Profile Analysis</p>
</div>

<div class="body-wrap">

<!-- META -->
<div class="meta">
    <span><b>File:</b> {_esc(file_label)}</span>
    <span><b>Date:</b> {now}</span>
    <span><b>Stage:</b> {_esc(a.get('career_stage','N/A'))}</span>
    <span><b>Score:</b> <span style="color:{score_color};font-weight:900;">{score}/100 ({score_label})</span></span>
</div>

<!-- CAREER SUMMARY -->
<div class="sec">
    <div class="sec-head">Career Summary</div>
    <p style="font-size:11px;color:#334155;line-height:1.7;margin:0 0 0 14px;">{_esc(a.get('career_summary','No summary available.'))}</p>
</div>

<!-- OPPORTUNITY SCORE -->
<div class="sec">
    <div class="sec-head">Opportunity Score</div>
    <div class="score-wrap">
        <div class="score-box">{score}</div>
        <p class="score-label">{score_label} Market Position</p>
    </div>
</div>

<!-- TECHNICAL SKILLS -->
<div class="sec">
    <div class="sec-head">Technical Skills</div>
    <div style="margin-left:14px;">{_pills(a.get('technical_skills',[]), '#4f46e5', '#eef2ff')}</div>
</div>

{"" if not a.get('tools_and_frameworks') else f'''
<div class="sec">
    <div class="sec-head">Tools &amp; Frameworks</div>
    <div style="margin-left:14px;">{_pills(a.get('tools_and_frameworks',[]), '#0e7490', '#ecfdf5')}</div>
</div>
'''}

<!-- SOFT SKILLS -->
<div class="sec">
    <div class="sec-head">Soft Skills</div>
    <div style="margin-left:14px;">{_pills(a.get('soft_skills',[]), '#7c3aed', '#f5f3ff')}</div>
</div>

{"" if not a.get('missing_skills') else f'''
<div class="sec">
    <div class="sec-head">Critical Missing Skills</div>
    <div style="margin-left:14px;">{_pills(a.get("missing_skills",[]), "#dc2626", "#fef2f2")}</div>
    <p style="font-size:9px;color:#94a3b8;margin:6px 0 0 14px;font-style:italic;">These skills are in high demand for your profile.</p>
</div>
'''}

<!-- JOB ROLES -->
<div class="sec">
    <div class="sec-head">Best-Fit Job Roles</div>
    <table>
        <tr><th>Role</th><th style="text-align:center;">Match</th><th>Reason</th></tr>
        {roles_html or '<tr><td colspan="3" class="empty">No roles detected</td></tr>'}
    </table>
</div>

<!-- INDUSTRY FIT -->
<div class="sec">
    <div class="sec-head">Industry Fit Analysis</div>
    <table>
        <tr><th>Industry</th><th style="text-align:center;">Fit %</th></tr>
        {ind_html or '<tr><td colspan="2" class="empty">No data</td></tr>'}
    </table>
</div>

<!-- SALARY -->
<div class="sec">
    <div class="sec-head">Salary Estimation (India)</div>
    <div class="sal-box">
        <span class="sal-min">{_format_salary(sal_min_raw)}</span>
        <span class="sal-arrow"> &#8594; </span>
        <span class="sal-max">{_format_salary(sal_max_raw)}</span>
        <span style="font-size:10px;color:#94a3b8;margin-left:10px;">per annum</span>
    </div>
    <p style="font-size:9px;color:#94a3b8;margin:2px 0 0 14px;">Based on skills, experience, and current Indian market trends.</p>
</div>

<!-- WEAKNESSES -->
<div class="sec">
    <div class="sec-head">Weaknesses Identified</div>
    {_bullets(a.get('weaknesses',[]), '&#10007;', '#dc2626')}
</div>

{"" if not a.get('certification_recommendations') else f'''
<div class="sec">
    <div class="sec-head">Recommended Certifications</div>
    {certs_html}
</div>
'''}

<!-- 3-MONTH PLAN -->
<div class="sec">
    <div class="sec-head">3-Month Improvement Plan</div>
    {_numbered(a.get('three_month_plan',[]), '#6366f1')}
</div>

<!-- 6-MONTH PLAN -->
<div class="sec">
    <div class="sec-head">6-Month Growth Roadmap</div>
    {_numbered(a.get('six_month_plan',[]), '#059669')}
</div>

<!-- ATS FEEDBACK -->
<div class="sec">
    <div class="sec-head">ATS Compatibility Feedback</div>
    {_bullets(a.get('ats_improvements',[]), '&#8594;', '#0e7490')}
</div>

{"" if not a.get('structure_issues') else f'''
<div class="sec">
    <div class="sec-head">Resume Structure Issues</div>
    {_bullets(a.get("structure_issues",[]), "&#9888;", "#d97706")}
</div>
'''}

{"" if not a.get('action_plan') else f'''
<div class="sec">
    <div class="sec-head">Personalized Action Plan</div>
    {_numbered(a.get("action_plan",[]), "#7c3aed")}
</div>
'''}

</div>

<!-- FOOTER -->
<div id="page_footer" style="background:#f8fafc;border-top:2px solid #e2e8f0;padding:8px 24px;text-align:center;">
    <span style="font-size:9px;color:#6366f1;font-weight:800;letter-spacing:1.5px;">RESUMATE AI</span>
    <span style="font-size:8px;color:#94a3b8;margin:0 6px;">&#8226;</span>
    <span style="font-size:8px;color:#94a3b8;">Career Intelligence Engine</span>
    <span style="font-size:8px;color:#94a3b8;margin:0 6px;">&#8226;</span>
    <span style="font-size:8px;color:#94a3b8;">{now} at {time_str}</span>
</div>

</body></html>"""

    buffer = io.BytesIO()
    pisa_status = pisa.CreatePDF(io.StringIO(html), dest=buffer)
    if pisa_status.err:
        raise Exception(f"PDF generation failed with {pisa_status.err} errors")
    buffer.seek(0)
    return buffer


# ═══════════════════════════════════════════════════════════════════════
#  WORD (.docx) REPORT GENERATION
# ═══════════════════════════════════════════════════════════════════════

def generate_docx_report(analysis_data, filename=None):
    """Generate a premium branded Word (.docx) career intelligence report."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    if isinstance(analysis_data, str):
        analysis_data = json.loads(analysis_data)

    a = analysis_data
    now = datetime.now().strftime("%B %d, %Y")
    time_str = datetime.now().strftime("%I:%M %p")
    file_label = filename or "Resume"
    score = a.get("opportunity_score", 0) or 0

    doc = Document()

    # ── Page Setup ──
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ── Base Style ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.line_spacing = 1.15

    # ── Color Palette ──
    BRAND    = RGBColor(0x31, 0x2e, 0x81)   # Deep indigo
    ACCENT   = RGBColor(0x63, 0x66, 0xf1)   # Indigo
    MUTED    = RGBColor(0x94, 0xa3, 0xb8)   # Slate-400
    DARK     = RGBColor(0x1e, 0x29, 0x3b)   # Slate-900
    MEDIUM   = RGBColor(0x47, 0x55, 0x69)   # Slate-600
    LIGHT    = RGBColor(0x64, 0x74, 0x8b)   # Slate-500
    GREEN    = RGBColor(0x05, 0x96, 0x69)
    AMBER    = RGBColor(0xd9, 0x77, 0x06)
    RED      = RGBColor(0xdc, 0x26, 0x26)
    CYAN     = RGBColor(0x08, 0x91, 0xb2)
    PURPLE   = RGBColor(0x7c, 0x3a, 0xed)
    DIVIDER  = RGBColor(0xe2, 0xe8, 0xf0)

    def score_color():
        if score >= 80:
            return GREEN
        if score >= 60:
            return AMBER
        return RED

    def format_salary(val):
        if not val:
            return "--"
        num = int(float(val))
        if num >= 100000:
            return f"\u20B9{num / 100000:.1f}L"
        if num >= 1000:
            return f"\u20B9{num // 1000}K"
        return f"\u20B9{num}"

    # ── Helper: shade table cells ──
    def shade_cells(table, row_idx, color_hex):
        for cell in table.rows[row_idx].cells:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
            cell._element.get_or_add_tcPr().append(shading)

    # ══════════════════════════════════════════════════════════
    #  TITLE PAGE
    # ══════════════════════════════════════════════════════════

    for _ in range(3):
        doc.add_paragraph()

    # Brand mark
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = brand.add_run("RESUMATE AI")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    r.font.name = 'Calibri'
    brand.paragraph_format.space_after = Pt(6)

    # Divider
    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = div.add_run("\u2500" * 30)
    r.font.color.rgb = DIVIDER
    r.font.size = Pt(10)

    doc.add_paragraph()

    # Main title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Career Intelligence\nReport")
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = BRAND

    doc.add_paragraph()

    # Tagline
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tag.add_run("Deep AI-Powered Career Profile Analysis")
    r.font.size = Pt(12)
    r.font.italic = True
    r.font.color.rgb = MUTED

    for _ in range(3):
        doc.add_paragraph()

    # Meta info
    meta_items = [
        f"File: {file_label}",
        f"Date: {now} at {time_str}",
        f"Career Stage: {a.get('career_stage', 'N/A')}",
        f"Opportunity Score: {score}/100",
    ]
    for text in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.color.rgb = LIGHT

    doc.add_paragraph()

    # Footer divider
    div2 = doc.add_paragraph()
    div2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = div2.add_run("\u2500" * 30)
    r.font.color.rgb = DIVIDER
    r.font.size = Pt(10)

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = foot.add_run("Confidential \u2022 For personal career development use")
    r.font.size = Pt(8)
    r.font.italic = True
    r.font.color.rgb = MUTED

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    #  HELPER FUNCTIONS
    # ══════════════════════════════════════════════════════════

    def add_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(4)
        bar = p.add_run("\u2503 ")
        bar.font.size = Pt(13)
        bar.font.bold = True
        bar.font.color.rgb = ACCENT
        r = p.add_run(text)
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = BRAND
        ul = doc.add_paragraph()
        ul.paragraph_format.space_after = Pt(6)
        r = ul.add_run("\u2500" * 55)
        r.font.color.rgb = DIVIDER
        r.font.size = Pt(7)

    def add_body(text, color=None):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(str(text))
        r.font.size = Pt(10.5)
        r.font.color.rgb = color or DARK

    def add_empty():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        r = p.add_run("None detected")
        r.font.italic = True
        r.font.color.rgb = MUTED
        r.font.size = Pt(10)

    def add_bullets(items, bullet="\u2022", color=None):
        if not items:
            add_empty()
            return
        for item in items:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.6)
            br = p.add_run(f"{bullet} ")
            br.font.bold = True
            br.font.color.rgb = color or ACCENT
            br.font.size = Pt(10.5)
            r = p.add_run(str(item))
            r.font.size = Pt(10.5)
            r.font.color.rgb = DARK

    def add_numbered(items, color=None):
        if not items:
            add_empty()
            return
        for i, item in enumerate(items, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.6)
            nr = p.add_run(f"{i}. ")
            nr.font.bold = True
            nr.font.color.rgb = color or ACCENT
            nr.font.size = Pt(10.5)
            r = p.add_run(str(item))
            r.font.size = Pt(10.5)
            r.font.color.rgb = DARK

    def add_skills_line(items, color=None):
        if not items:
            add_empty()
            return
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(6)
        for i, item in enumerate(items):
            r = p.add_run(f" {item} ")
            r.font.size = Pt(10)
            r.font.bold = True
            r.font.color.rgb = color or ACCENT
            if i < len(items) - 1:
                s = p.add_run("  |  ")
                s.font.color.rgb = DIVIDER
                s.font.size = Pt(10)

    def add_table_styled(headers, rows):
        """Add a styled table with header shading."""
        t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(headers):
            cell = t.rows[0].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(h)
            r.font.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = MEDIUM
        shade_cells(t, 0, "F1F5F9")
        for i, row_data in enumerate(rows, 1):
            for j, val in enumerate(row_data):
                cell = t.rows[i].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(str(val))
                r.font.size = Pt(10)
            if i % 2 == 0:
                shade_cells(t, i, "FAFAFA")
        return t

    # ══════════════════════════════════════════════════════════
    #  REPORT CONTENT
    # ══════════════════════════════════════════════════════════

    # ── EXECUTIVE SUMMARY ──
    add_heading("Executive Summary")

    sal = a.get('salary_estimation_inr', {})
    metrics = [
        ("Career Stage", a.get("career_stage", "N/A")),
        ("Opportunity Score", f"{score}/100"),
        ("Salary Range", f"{format_salary(sal.get('min'))} \u2013 {format_salary(sal.get('max'))} p.a."),
        ("Technical Skills", str(len(a.get("technical_skills", [])))),
        ("Missing Skills", str(len(a.get("missing_skills", [])))),
        ("Top Job Match", (a["job_roles"][0]["role"] + f' ({a["job_roles"][0]["match_percentage"]}%)') if a.get("job_roles") else "N/A"),
    ]
    t = doc.add_table(rows=len(metrics), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(metrics):
        t.rows[i].cells[0].text = ""
        p = t.rows[i].cells[0].paragraphs[0]
        r = p.add_run(label)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = MEDIUM
        t.rows[i].cells[1].text = ""
        p = t.rows[i].cells[1].paragraphs[0]
        r = p.add_run(value)
        r.font.size = Pt(10)
        if i == 1:
            r.font.bold = True
            r.font.color.rgb = score_color()
        if i % 2 == 0:
            shade_cells(t, i, "F8FAFC")

    doc.add_paragraph()
    add_body(a.get('career_summary', 'No summary available.'))

    # ── TECHNICAL SKILLS ──
    add_heading("Technical Skills")
    add_skills_line(a.get('technical_skills', []), ACCENT)

    if a.get('tools_and_frameworks'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.left_indent = Cm(0.4)
        r = p.add_run("Tools & Frameworks:")
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = MEDIUM
        add_skills_line(a.get('tools_and_frameworks', []), CYAN)

    # ── SOFT SKILLS ──
    add_heading("Soft Skills")
    add_skills_line(a.get('soft_skills', []), PURPLE)

    # ── MISSING SKILLS ──
    if a.get('missing_skills'):
        add_heading("Critical Missing Skills")
        add_skills_line(a.get('missing_skills', []), RED)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        r = p.add_run("These skills are in high demand for your profile \u2014 consider prioritizing them.")
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = MUTED

    # ── JOB ROLES ──
    add_heading("Best-Fit Job Roles")
    roles = a.get('job_roles', [])
    if roles:
        rows = [(jr.get('role', ''), f"{jr.get('match_percentage', 0)}%", jr.get('reason', '')) for jr in roles]
        add_table_styled(["Role", "Match %", "Reason"], rows)
    else:
        add_empty()

    # ── INDUSTRY FIT ──
    add_heading("Industry Fit Analysis")
    industries = a.get('industry_fit', [])
    if industries:
        rows = [(ind.get('industry', ''), f"{ind.get('fit_percentage', 0)}%") for ind in industries]
        add_table_styled(["Industry", "Fit %"], rows)
    else:
        add_empty()

    # ── SALARY ──
    add_heading("Salary Estimation (India)")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"Minimum: {format_salary(sal.get('min'))}")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = AMBER
    r = p.add_run("    \u2192    ")
    r.font.size = Pt(14)
    r.font.color.rgb = DIVIDER
    r = p.add_run(f"Maximum: {format_salary(sal.get('max'))}")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = GREEN
    note = doc.add_paragraph()
    note.paragraph_format.left_indent = Cm(0.4)
    r = note.add_run("Based on skills, experience level, and current Indian job market data.")
    r.font.size = Pt(8.5)
    r.font.italic = True
    r.font.color.rgb = MUTED

    # ── WEAKNESSES ──
    add_heading("Weaknesses Identified")
    add_bullets(a.get('weaknesses', []), "\u2717", RED)

    # ── CERTIFICATIONS ──
    if a.get('certification_recommendations'):
        add_heading("Recommended Certifications")
        for cert in a['certification_recommendations']:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(3)
            if isinstance(cert, str):
                r = p.add_run(f"\u2022 {cert}")
                r.font.size = Pt(10.5)
            else:
                r = p.add_run(f"\u2022 {cert.get('name', '')}")
                r.font.size = Pt(10.5)
                r.font.bold = True
                if cert.get('provider'):
                    r = p.add_run(f"  ({cert['provider']})")
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = ACCENT
                if cert.get('reason'):
                    r = p.add_run(f" \u2014 {cert['reason']}")
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = LIGHT

    # ── 3-MONTH PLAN ──
    add_heading("3-Month Improvement Plan")
    add_numbered(a.get('three_month_plan', []), ACCENT)

    # ── 6-MONTH PLAN ──
    add_heading("6-Month Growth Roadmap")
    add_numbered(a.get('six_month_plan', []), GREEN)

    # ── ATS FEEDBACK ──
    add_heading("ATS Compatibility Feedback")
    add_bullets(a.get('ats_improvements', []), "\u2192", CYAN)

    # ── STRUCTURE ISSUES ──
    if a.get('structure_issues'):
        add_heading("Resume Structure Issues")
        add_bullets(a.get('structure_issues', []), "\u26A0", AMBER)

    # ── ACTION PLAN ──
    if a.get('action_plan'):
        add_heading("Personalized Action Plan")
        add_numbered(a.get('action_plan', []), PURPLE)

    # ══════════════════════════════════════════════════════════
    #  REPORT FOOTER
    # ══════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    div_f = doc.add_paragraph()
    div_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = div_f.add_run("\u2500" * 40)
    r.font.color.rgb = DIVIDER
    r.font.size = Pt(7)

    fb = doc.add_paragraph()
    fb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fb.paragraph_format.space_after = Pt(2)
    r = fb.add_run("RESUMATE AI")
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = ACCENT

    fb2 = doc.add_paragraph()
    fb2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fb2.paragraph_format.space_after = Pt(2)
    r = fb2.add_run("Career Intelligence Engine")
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED

    fb3 = doc.add_paragraph()
    fb3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fb3.add_run(f"Report generated on {now} at {time_str}")
    r.font.size = Pt(8)
    r.font.italic = True
    r.font.color.rgb = MUTED

    fb4 = doc.add_paragraph()
    fb4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fb4.add_run("This report is AI-generated. Please review before making career decisions.")
    r.font.size = Pt(7.5)
    r.font.italic = True
    r.font.color.rgb = MUTED

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ═══════════════════════════════════════════════════════════════════════
#  COMPARISON PDF REPORT
# ═══════════════════════════════════════════════════════════════════════

def generate_comparison_pdf(comparison_data):
    """Generate a branded PDF comparison report from multi-resume analysis."""
    from xhtml2pdf import pisa

    if isinstance(comparison_data, str):
        comparison_data = json.loads(comparison_data)

    c = comparison_data
    now = datetime.now().strftime("%B %d, %Y")
    time_str = datetime.now().strftime("%I:%M %p")
    target = c.get("target_job_profile", "General Career Strength")
    best = c.get("best_resume", {})

    def _esc(t):
        if not t: return ""
        return str(t).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    # Ranking table rows
    ranking_rows = ""
    for r in c.get("ranking", []):
        sc = r.get("match_score", 0)
        clr = "#059669" if sc >= 80 else "#e67e22" if sc >= 60 else "#e74c3c"
        medal = {1: "🥇", 2: "🥈", 3: "🥉"}.get(r.get("rank", 99), f"#{r.get('rank','')}")
        ranking_rows += f'''<tr>
            <td style="padding:8px 12px;border-bottom:1px solid #f1f5f9;font-size:12px;font-weight:800;text-align:center;">{medal}</td>
            <td style="padding:8px 12px;border-bottom:1px solid #f1f5f9;font-size:11px;font-weight:700;">{_esc(r.get("resume_name",""))}</td>
            <td style="padding:8px 12px;border-bottom:1px solid #f1f5f9;font-size:12px;font-weight:900;color:{clr};text-align:center;">{sc}</td>
            <td style="padding:8px 12px;border-bottom:1px solid #f1f5f9;font-size:10px;color:#64748b;">{_esc(r.get("one_line_verdict",""))}</td>
        </tr>'''

    # Individual results
    results_html = ""
    for res in c.get("comparison_results", []):
        sc = res.get("match_score", 0)
        clr = "#059669" if sc >= 80 else "#e67e22" if sc >= 60 else "#e74c3c"
        is_best = res.get("resume_name", "") == best.get("resume_name", "")
        badge = '<span style="background:#059669;color:white;padding:2px 8px;border-radius:8px;font-size:9px;font-weight:800;margin-left:8px;">BEST MATCH</span>' if is_best else ''
        skills = ", ".join(res.get("key_skills", [])[:8])
        strengths = "".join(f'<p style="margin:2px 0;font-size:10px;color:#059669;">✓ {_esc(s)}</p>' for s in res.get("strengths", []))
        weaknesses = "".join(f'<p style="margin:2px 0;font-size:10px;color:#dc2626;">✗ {_esc(w)}</p>' for w in res.get("weaknesses", []))
        suggestions = "".join(f'<p style="margin:2px 0;font-size:10px;color:#6366f1;">→ {_esc(s)}</p>' for s in res.get("improvement_suggestions", []))

        results_html += f'''
        <div style="border:1px solid {"#059669" if is_best else "#e2e8f0"};border-radius:8px;padding:14px 18px;margin-bottom:14px;{"background:#f0fdf4;" if is_best else ""}page-break-inside:avoid;">
            <p style="font-size:13px;font-weight:800;color:#1e293b;margin:0 0 4px 0;">{_esc(res.get("resume_name",""))}{badge}
                <span style="float:right;font-size:16px;font-weight:900;color:{clr};">{sc}/100</span>
            </p>
            <p style="font-size:9px;color:#94a3b8;margin:0 0 8px 0;">Stage: {_esc(res.get("career_stage","N/A"))} | Skills: {_esc(skills)}</p>
            <table style="width:100%;"><tr>
                <td style="vertical-align:top;width:33%;padding-right:8px;"><p style="font-size:9px;font-weight:700;color:#475569;margin:0 0 4px 0;">STRENGTHS</p>{strengths or "<p style='color:#94a3b8;font-size:9px;'>—</p>"}</td>
                <td style="vertical-align:top;width:33%;padding-right:8px;"><p style="font-size:9px;font-weight:700;color:#475569;margin:0 0 4px 0;">WEAKNESSES</p>{weaknesses or "<p style='color:#94a3b8;font-size:9px;'>—</p>"}</td>
                <td style="vertical-align:top;width:34%;"><p style="font-size:9px;font-weight:700;color:#475569;margin:0 0 4px 0;">SUGGESTIONS</p>{suggestions or "<p style='color:#94a3b8;font-size:9px;'>—</p>"}</td>
            </tr></table>
        </div>'''

    # Best resume details
    details = best.get("details", {})
    best_details_html = ""
    for key, label in [("skills_advantage", "Skills"), ("experience_advantage", "Experience"),
                       ("education_advantage", "Education"), ("achievements_advantage", "Achievements"),
                       ("ats_advantage", "ATS Readiness")]:
        val = details.get(key, "")
        if val:
            best_details_html += f'<p style="margin:4px 0;font-size:11px;"><b style="color:#312e81;">{label}:</b> {_esc(val)}</p>'

    html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"/>
<style>
@page {{ size: A4; margin: 1.2cm 1cm 2cm 1.5cm;
    @frame footer_frame {{ -pdf-frame-content: page_footer; bottom: 0; left: 0; right: 0; height: 1.8cm; }}
}}
body {{ font-family: Helvetica, Arial, sans-serif; color: #1e293b; font-size: 11px; line-height: 1.5; margin: 0; padding: 0; }}
.header {{ background-color: #312e81; padding: 28px 24px 22px 24px; color: white; }}
.header-brand {{ font-size: 11px; font-weight: 800; letter-spacing: 3px; text-transform: uppercase; color: #a5b4fc; margin: 0 0 2px 0; }}
.header-title {{ font-size: 24px; font-weight: 800; margin: 0 0 4px 0; }}
.header-sub {{ font-size: 10px; color: #c7d2fe; margin: 0; }}
.body-wrap {{ padding: 20px 18px 20px 22px; }}
.sec {{ margin-bottom: 16px; page-break-inside: avoid; }}
.sec-head {{ font-size: 13px; font-weight: 800; color: #312e81; border-left: 3px solid #6366f1; padding-left: 10px; margin-bottom: 10px; }}
table {{ width: 100%; border-collapse: collapse; }}
th {{ background: #f1f5f9; color: #475569; font-size: 9px; font-weight: 800; text-transform: uppercase; padding: 7px 12px; text-align: left; border-bottom: 2px solid #cbd5e1; }}
</style></head>
<body>

<div class="header">
    <p class="header-brand">ResuMate AI</p>
    <p class="header-title">Resume Comparison Report</p>
    <p class="header-sub">Multi-Resume AI Analysis &amp; Best Fit Selection</p>
</div>

<div class="body-wrap">

<div style="background:#f8fafc;border:1px solid #e2e8f0;border-radius:6px;padding:10px 16px;margin-bottom:18px;">
    <span style="font-size:10px;color:#64748b;margin-right:18px;"><b>Date:</b> {now}</span>
    <span style="font-size:10px;color:#64748b;margin-right:18px;"><b>Target:</b> {_esc(target)}</span>
    <span style="font-size:10px;color:#64748b;"><b>Resumes:</b> {c.get("total_resumes", len(c.get("comparison_results", [])))}</span>
</div>

<div class="sec">
    <div class="sec-head">Final Ranking</div>
    <table>
        <tr><th style="text-align:center;">Rank</th><th>Resume</th><th style="text-align:center;">Score</th><th>Verdict</th></tr>
        {ranking_rows}
    </table>
</div>

<div class="sec">
    <div class="sec-head">Winner: {_esc(best.get("resume_name",""))}</div>
    <div style="background:#f0fdf4;border:2px solid #059669;border-radius:8px;padding:14px 18px;">
        <p style="font-size:12px;font-weight:700;color:#059669;margin:0 0 6px 0;">Why This Resume Wins</p>
        <p style="font-size:11px;color:#334155;margin:0 0 10px 0;">{_esc(best.get("reason",""))}</p>
        {best_details_html}
    </div>
</div>

<div class="sec">
    <div class="sec-head">Individual Resume Analysis</div>
    {results_html}
</div>

<div class="sec">
    <div class="sec-head">Comparison Summary</div>
    <p style="font-size:11px;color:#334155;line-height:1.7;margin:0 0 0 14px;">{_esc(c.get("comparison_summary",""))}</p>
</div>

</div>

<div id="page_footer" style="background:#f8fafc;border-top:2px solid #e2e8f0;padding:8px 24px;text-align:center;">
    <span style="font-size:9px;color:#6366f1;font-weight:800;letter-spacing:1.5px;">RESUMATE AI</span>
    <span style="font-size:8px;color:#94a3b8;margin:0 6px;">&#8226;</span>
    <span style="font-size:8px;color:#94a3b8;">Resume Comparison Engine</span>
    <span style="font-size:8px;color:#94a3b8;margin:0 6px;">&#8226;</span>
    <span style="font-size:8px;color:#94a3b8;">{now} at {time_str}</span>
</div>

</body></html>"""

    buffer = io.BytesIO()
    pisa_status = pisa.CreatePDF(io.StringIO(html), dest=buffer)
    if pisa_status.err:
        raise Exception(f"PDF generation failed with {pisa_status.err} errors")
    buffer.seek(0)
    return buffer


# ═══════════════════════════════════════════════════════════════════════
#  COMPARISON DOCX REPORT
# ═══════════════════════════════════════════════════════════════════════

def generate_comparison_docx(comparison_data):
    """Generate a branded Word comparison report."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    if isinstance(comparison_data, str):
        comparison_data = json.loads(comparison_data)

    c = comparison_data
    now = datetime.now().strftime("%B %d, %Y")
    time_str = datetime.now().strftime("%I:%M %p")
    target = c.get("target_job_profile", "General Career Strength")
    best = c.get("best_resume", {})

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

    BRAND   = RGBColor(0x31, 0x2e, 0x81)
    ACCENT  = RGBColor(0x63, 0x66, 0xf1)
    MUTED   = RGBColor(0x94, 0xa3, 0xb8)
    DARK    = RGBColor(0x1e, 0x29, 0x3b)
    MEDIUM  = RGBColor(0x47, 0x55, 0x69)
    GREEN   = RGBColor(0x05, 0x96, 0x69)
    RED     = RGBColor(0xdc, 0x26, 0x26)
    AMBER   = RGBColor(0xd9, 0x77, 0x06)
    DIVIDER = RGBColor(0xe2, 0xe8, 0xf0)

    def shade_cells(table, row_idx, color_hex):
        for cell in table.rows[row_idx].cells:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
            cell._element.get_or_add_tcPr().append(shading)

    def add_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(4)
        bar = p.add_run("\u2503 ")
        bar.font.size = Pt(13)
        bar.font.bold = True
        bar.font.color.rgb = ACCENT
        r = p.add_run(text)
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = BRAND
        ul = doc.add_paragraph()
        ul.paragraph_format.space_after = Pt(6)
        r = ul.add_run("\u2500" * 55)
        r.font.color.rgb = DIVIDER
        r.font.size = Pt(7)

    # ── TITLE PAGE ──
    for _ in range(3):
        doc.add_paragraph()

    b = doc.add_paragraph()
    b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = b.add_run("RESUMATE AI")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = ACCENT

    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = div.add_run("\u2500" * 30)
    r.font.color.rgb = DIVIDER

    doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Resume Comparison\nReport")
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = BRAND

    doc.add_paragraph()

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tag.add_run("Multi-Resume AI Analysis & Best Fit Selection")
    r.font.size = Pt(12)
    r.font.italic = True
    r.font.color.rgb = MUTED

    for _ in range(3):
        doc.add_paragraph()

    for text in [f"Target: {target}", f"Resumes Compared: {c.get('total_resumes', 0)}", f"Date: {now} at {time_str}"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.color.rgb = MEDIUM

    doc.add_page_break()

    # ── RANKING TABLE ──
    add_heading("Final Ranking")
    ranking = c.get("ranking", [])
    if ranking:
        tbl = doc.add_table(rows=len(ranking) + 1, cols=4)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(["Rank", "Resume", "Score", "Verdict"]):
            tbl.rows[0].cells[j].text = ""
            p = tbl.rows[0].cells[j].paragraphs[0]
            r = p.add_run(h)
            r.font.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = MEDIUM
        shade_cells(tbl, 0, "F1F5F9")
        for i, rk in enumerate(ranking, 1):
            sc = rk.get("match_score", 0)
            sc_color = GREEN if sc >= 80 else AMBER if sc >= 60 else RED
            medal = {1: "\U0001F947", 2: "\U0001F948", 3: "\U0001F949"}.get(rk.get("rank", 99), f"#{rk.get('rank','')}")
            for j, val in enumerate([medal, rk.get("resume_name", ""), str(sc), rk.get("one_line_verdict", "")]):
                tbl.rows[i].cells[j].text = ""
                p = tbl.rows[i].cells[j].paragraphs[0]
                r = p.add_run(str(val))
                r.font.size = Pt(10)
                if j == 2:
                    r.font.bold = True
                    r.font.color.rgb = sc_color
            if i % 2 == 0:
                shade_cells(tbl, i, "FAFAFA")

    # ── BEST RESUME ──
    add_heading(f"Winner: {best.get('resume_name', '')}")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run("Why This Resume Wins")
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = GREEN

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    r = p.add_run(best.get("reason", ""))
    r.font.size = Pt(10.5)

    details = best.get("details", {})
    for key, label in [("skills_advantage", "Skills"), ("experience_advantage", "Experience"),
                       ("education_advantage", "Education"), ("achievements_advantage", "Achievements"),
                       ("ats_advantage", "ATS Readiness")]:
        val = details.get(key, "")
        if val:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(f"{label}: ")
            r.font.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = BRAND
            r = p.add_run(val)
            r.font.size = Pt(10)

    # ── INDIVIDUAL RESULTS ──
    add_heading("Individual Resume Analysis")
    for res in c.get("comparison_results", []):
        is_best = res.get("resume_name", "") == best.get("resume_name", "")
        sc = res.get("match_score", 0)
        sc_color = GREEN if sc >= 80 else AMBER if sc >= 60 else RED

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        r = p.add_run(f"{res.get('resume_name', '')}  —  {sc}/100")
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = GREEN if is_best else DARK
        if is_best:
            r = p.add_run("  [BEST MATCH]")
            r.font.size = Pt(9)
            r.font.bold = True
            r.font.color.rgb = GREEN

        for section_name, items, bullet, color in [
            ("Strengths", res.get("strengths", []), "\u2713", GREEN),
            ("Weaknesses", res.get("weaknesses", []), "\u2717", RED),
            ("Suggestions", res.get("improvement_suggestions", []), "\u2192", ACCENT),
        ]:
            if items:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.4)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(section_name)
                r.font.bold = True
                r.font.size = Pt(9.5)
                r.font.color.rgb = MEDIUM
                for item in items:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.8)
                    p.paragraph_format.space_after = Pt(1)
                    r = p.add_run(f"{bullet} ")
                    r.font.color.rgb = color
                    r.font.size = Pt(10)
                    r = p.add_run(str(item))
                    r.font.size = Pt(10)

    # ── COMPARISON SUMMARY ──
    summary = c.get("comparison_summary", "")
    if summary:
        add_heading("Comparison Summary")
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        r = p.add_run(summary)
        r.font.size = Pt(10.5)

    # ── FOOTER ──
    doc.add_paragraph()
    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = div.add_run("\u2500" * 40)
    r.font.color.rgb = DIVIDER
    r.font.size = Pt(7)

    for text, size, color, bold, italic in [
        ("RESUMATE AI", 9, ACCENT, True, False),
        ("Resume Comparison Engine", 8, MUTED, False, False),
        (f"Report generated on {now} at {time_str}", 8, MUTED, False, True),
        ("This report is AI-generated. Please review before making hiring decisions.", 7.5, MUTED, False, True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.font.bold = bold
        r.font.italic = italic

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ═══════════════════════════════════════════════════════════════════════
#  JOB FINDER REPORT — PDF (xhtml2pdf)
# ═══════════════════════════════════════════════════════════════════════

def generate_jobs_pdf(data):
    """Generate a branded PDF Job Intelligence Report."""
    from xhtml2pdf import pisa

    now = datetime.now().strftime("%B %d, %Y")
    time_str = datetime.now().strftime("%I:%M %p")

    def _esc(text):
        if not text:
            return ""
        return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    rec_jobs = data.get("recommended_jobs", [])
    all_jobs_list = data.get("all_jobs", [])

    rec_rows = ""
    for group in rec_jobs:
        for job in group.get("jobs", []):
            score = job.get("match_score", "—")
            rec_rows += (
                '<tr>'
                f'<td style="padding:8px 10px;border-bottom:1px solid #f1f5f9;font-size:10px;">{_esc(job.get("title",""))}</td>'
                f'<td style="padding:8px 10px;border-bottom:1px solid #f1f5f9;font-size:10px;">{_esc(job.get("company",""))}</td>'
                f'<td style="padding:8px 10px;border-bottom:1px solid #f1f5f9;font-size:10px;">{_esc(job.get("location",""))}</td>'
                f'<td style="padding:8px 10px;border-bottom:1px solid #f1f5f9;font-size:10px;">{_esc(job.get("salary_range","N/A"))}</td>'
                f'<td style="padding:8px 10px;border-bottom:1px solid #f1f5f9;font-size:10px;font-weight:bold;color:#059669;">{score}%</td>'
                '</tr>'
            )

    all_rows = ""
    for job in all_jobs_list[:50]:
        all_rows += (
            '<tr>'
            f'<td style="padding:8px 10px;border-bottom:1px solid #f1f5f9;font-size:10px;">{_esc(job.get("title",""))}</td>'
            f'<td style="padding:8px 10px;border-bottom:1px solid #f1f5f9;font-size:10px;">{_esc(job.get("company",""))}</td>'
            f'<td style="padding:8px 10px;border-bottom:1px solid #f1f5f9;font-size:10px;">{_esc(job.get("location",""))}</td>'
            f'<td style="padding:8px 10px;border-bottom:1px solid #f1f5f9;font-size:10px;">{_esc(job.get("salary_range","N/A"))}</td>'
            '</tr>'
        )

    html = f"""<!DOCTYPE html>
    <html><head><meta charset="utf-8">
    <style>
        @page {{ size: A4; margin: 1.5cm; }}
        body {{ font-family: Helvetica, Arial, sans-serif; color: #1e293b; font-size: 11px; line-height: 1.6; }}
        .header {{ background: #059669; color: white; padding: 28px 24px; margin-bottom: 20px; }}
        .header h1 {{ font-size: 22px; margin: 0; font-weight: 800; }}
        .header p {{ margin: 4px 0 0; font-size: 10px; opacity: 0.85; }}
        .section {{ margin-bottom: 18px; border: 1px solid #e2e8f0; padding: 16px; }}
        .section-title {{ font-size: 14px; font-weight: 800; color: #059669; margin: 0 0 10px; }}
        table {{ width: 100%; border-collapse: collapse; }}
        th {{ background: #f0fdf4; padding: 8px 10px; font-size: 9px; text-transform: uppercase;
              font-weight: 700; color: #059669; text-align: left; border-bottom: 2px solid #a7f3d0; }}
        .footer {{ text-align: center; margin-top: 30px; padding-top: 14px;
                   border-top: 1px solid #e2e8f0; font-size: 9px; color: #94a3b8; }}
    </style></head><body>
    <div class="header">
        <h1>ResuMate AI Job Intelligence Report</h1>
        <p>Generated on {now} at {time_str}</p>
    </div>
    """

    if rec_rows:
        html += f"""
        <div class="section">
            <h2 class="section-title">Recommended Job Listings</h2>
            <table><tr><th>Job Title</th><th>Company</th><th>Location</th><th>Salary</th><th>Match</th></tr>
            {rec_rows}</table>
        </div>"""

    if all_rows:
        html += f"""
        <div class="section">
            <h2 class="section-title">Available Job Listings</h2>
            <table><tr><th>Job Title</th><th>Company</th><th>Location</th><th>Salary</th></tr>
            {all_rows}</table>
        </div>"""

    html += f"""
    <div class="footer">
        <p><strong>RESUMATE AI</strong> — Job Intelligence Engine</p>
        <p>Report generated on {now} at {time_str}</p>
        <p>Powered by Adzuna. Data is indicative and subject to change.</p>
    </div>
    </body></html>"""

    buffer = io.BytesIO()
    pisa.CreatePDF(io.StringIO(html), dest=buffer)
    buffer.seek(0)
    return buffer


# ═══════════════════════════════════════════════════════════════════════
#  JOB FINDER REPORT — DOCX (python-docx)
# ═══════════════════════════════════════════════════════════════════════

def generate_jobs_docx(data):
    """Generate a branded Word Job Intelligence Report."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    ACCENT = RGBColor(5, 150, 105)
    DARK = RGBColor(30, 41, 59)
    MUTED = RGBColor(100, 116, 139)

    now = datetime.now().strftime("%B %d, %Y")
    time_str = datetime.now().strftime("%I:%M %p")

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = DARK

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("RESUMATE AI")
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = ACCENT

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub_p.add_run("Job Intelligence Report")
    r.font.size = Pt(14)
    r.font.color.rgb = DARK
    r.font.bold = True

    datep = doc.add_paragraph()
    datep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = datep.add_run(f"Generated on {now} at {time_str}")
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED

    doc.add_paragraph()

    rec_jobs = data.get("recommended_jobs", [])
    all_jobs_list = data.get("all_jobs", [])

    if rec_jobs:
        sh = doc.add_paragraph()
        r = sh.add_run("RECOMMENDED JOB LISTINGS")
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = ACCENT

        flat_jobs = []
        for group in rec_jobs:
            for job in group.get("jobs", []):
                flat_jobs.append(job)

        if flat_jobs:
            table = doc.add_table(rows=1, cols=5)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            for idx, label in enumerate(["Job Title", "Company", "Location", "Salary", "Match %"]):
                hdr[idx].text = label
                for para in hdr[idx].paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = ACCENT

            for job in flat_jobs[:30]:
                row = table.add_row().cells
                row[0].text = job.get("title", "")
                row[1].text = job.get("company", "")
                row[2].text = job.get("location", "")
                row[3].text = job.get("salary_range", "N/A")
                row[4].text = f"{job.get('match_score', '—')}%"

            doc.add_paragraph()

    if all_jobs_list:
        sh = doc.add_paragraph()
        r = sh.add_run("AVAILABLE JOB LISTINGS")
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = ACCENT

        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for idx, label in enumerate(["Job Title", "Company", "Location", "Salary"]):
            hdr[idx].text = label
            for para in hdr[idx].paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = ACCENT

        for job in all_jobs_list[:50]:
            row = table.add_row().cells
            row[0].text = job.get("title", "")
            row[1].text = job.get("company", "")
            row[2].text = job.get("location", "")
            row[3].text = job.get("salary_range", "N/A")

    doc.add_paragraph()
    for text, size, color, bold in [
        ("RESUMATE AI", 9, ACCENT, True),
        ("Job Intelligence Engine", 8, MUTED, False),
        (f"Report generated on {now} at {time_str}", 8, MUTED, False),
        ("Powered by Adzuna. Data is indicative and subject to change.", 7.5, MUTED, False),
    ]:
        fp = doc.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_after = Pt(2)
        r = fp.add_run(text)
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.font.bold = bold

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ═══════════════════════════════════════════════════════════════════════
#  COMPANY INSIGHTS REPORT — PDF (xhtml2pdf)
# ═══════════════════════════════════════════════════════════════════════

def generate_company_pdf(data):
    """Generate a branded PDF Company Intelligence Report."""
    from xhtml2pdf import pisa

    now = datetime.now().strftime("%B %d, %Y")
    time_str = datetime.now().strftime("%I:%M %p")
    company_name = data.get("company_name", "Company")
    job_role = data.get("job_role", "")
    a = data.get("analysis", {})

    def _esc(t):
        if not t:
            return ""
        return str(t).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def _csection(num, stitle, content):
        if not content:
            return ""
        return (
            '<div class="section">'
            f'<div class="section-head"><span class="num">{num}</span>'
            f'<span class="stitle">{_esc(stitle)}</span></div>'
            f'<p>{_esc(content)}</p>'
            '</div>'
        )

    skills_html = ""
    for s in (a.get("key_skills_they_value") or []):
        skills_html += f'<span class="skill">{_esc(s)}</span>'

    tips_html = ""
    for i, t in enumerate(a.get("preparation_tips") or []):
        tips_html += f'<li><strong>{i+1}.</strong> {_esc(t)}</li>'

    html = f"""<!DOCTYPE html>
    <html><head><meta charset="utf-8">
    <style>
        @page {{ size: A4; margin: 1.5cm; }}
        body {{ font-family: Helvetica, Arial, sans-serif; color: #1e293b; font-size: 11px; line-height: 1.6; }}
        .header {{ background: #6366f1; color: white; padding: 28px 24px; margin-bottom: 20px; }}
        .header h1 {{ font-size: 20px; margin: 0; font-weight: 800; }}
        .header p {{ margin: 4px 0 0; font-size: 10px; opacity: 0.85; }}
        .section {{ margin-bottom: 14px; border: 1px solid #e2e8f0; padding: 14px 16px; }}
        .section-head {{ margin-bottom: 8px; }}
        .num {{ display: inline-block; background: #eef2ff; color: #6366f1; padding: 2px 8px;
                font-size: 9px; font-weight: 800; margin-right: 8px; }}
        .stitle {{ font-size: 13px; font-weight: 800; color: #334155; }}
        p {{ font-size: 10.5px; line-height: 1.65; margin: 0; }}
        .skill {{ display: inline-block; background: #ecfdf5; color: #059669; padding: 3px 10px;
                  font-size: 9px; font-weight: 700; margin: 2px; }}
        ul {{ padding-left: 18px; margin: 4px 0; }}
        li {{ font-size: 10.5px; margin-bottom: 4px; }}
        .match {{ background: #ecfdf5; border: 1px solid #a7f3d0; padding: 14px; margin: 16px 0; }}
        .match h3 {{ color: #059669; font-size: 12px; margin: 0 0 6px; }}
        .footer {{ text-align: center; margin-top: 30px; padding-top: 14px;
                   border-top: 1px solid #e2e8f0; font-size: 9px; color: #94a3b8; }}
    </style></head><body>
    <div class="header">
        <h1>{_esc(company_name)} — Company Intelligence Report</h1>
        <p>{"Role: " + _esc(job_role) + " | " if job_role else ""}Generated on {now} at {time_str}</p>
    </div>
    {_csection("01", "Company Overview", a.get("company_overview"))}
    {_csection("02", "Industry Position", a.get("industry_position"))}
    {_csection("03", "Company Size", a.get("company_size_estimate"))}
    {_csection("04", "Work Culture", a.get("work_culture"))}
    {_csection("05", "Interview Process", a.get("interview_process"))}
    """

    if skills_html:
        html += (
            '<div class="section">'
            '<div class="section-head"><span class="num">06</span>'
            '<span class="stitle">Key Skills They Value</span></div>'
            f'{skills_html}</div>'
        )

    html += _csection("07", "Salary Expectation (INR)", a.get("salary_expectation_inr"))
    html += _csection("08", "Growth Opportunities", a.get("growth_opportunities"))
    html += _csection("09", "Risks & Challenges", a.get("risks"))

    if tips_html:
        html += (
            '<div class="section">'
            '<div class="section-head"><span class="num">10</span>'
            '<span class="stitle">Preparation Tips</span></div>'
            f'<ul>{tips_html}</ul></div>'
        )

    if a.get("profile_match"):
        html += (
            '<div class="match">'
            '<h3>Why You Match This Company</h3>'
            f'<p>{_esc(a.get("profile_match"))}</p></div>'
        )

    html += f"""
    <div class="footer">
        <p><strong>RESUMATE AI</strong> — Company Intelligence Engine</p>
        <p>Report generated on {now} at {time_str}</p>
    </div>
    </body></html>"""

    buffer = io.BytesIO()
    pisa.CreatePDF(io.StringIO(html), dest=buffer)
    buffer.seek(0)
    return buffer


# ═══════════════════════════════════════════════════════════════════════
#  COMPANY INSIGHTS REPORT — DOCX (python-docx)
# ═══════════════════════════════════════════════════════════════════════

def generate_company_docx(data):
    """Generate a branded Word Company Intelligence Report."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ACCENT = RGBColor(99, 102, 241)
    GREEN = RGBColor(5, 150, 105)
    DARK = RGBColor(30, 41, 59)
    MUTED = RGBColor(100, 116, 139)

    now = datetime.now().strftime("%B %d, %Y")
    time_str = datetime.now().strftime("%I:%M %p")
    company_name = data.get("company_name", "Company")
    job_role = data.get("job_role", "")
    a = data.get("analysis", {})

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = DARK

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(company_name.upper())
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = ACCENT

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub_p.add_run("Company Intelligence Report")
    r.font.size = Pt(14)
    r.font.color.rgb = DARK
    r.font.bold = True

    if job_role:
        rp = doc.add_paragraph()
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = rp.add_run(f"Role: {job_role}")
        r.font.size = Pt(10)
        r.font.color.rgb = MUTED

    datep = doc.add_paragraph()
    datep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = datep.add_run(f"Generated on {now} at {time_str}")
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED

    doc.add_paragraph()

    def add_sect(num, title_text, content):
        if not content:
            return
        h = doc.add_paragraph()
        r = h.add_run(f"{num}  {title_text}")
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = ACCENT
        p = doc.add_paragraph(str(content))
        p.paragraph_format.space_after = Pt(8)
        doc.add_paragraph()

    add_sect("01", "Company Overview", a.get("company_overview"))
    add_sect("02", "Industry Position", a.get("industry_position"))
    add_sect("03", "Company Size", a.get("company_size_estimate"))
    add_sect("04", "Work Culture", a.get("work_culture"))
    add_sect("05", "Interview Process", a.get("interview_process"))

    if a.get("key_skills_they_value"):
        h = doc.add_paragraph()
        r = h.add_run("06  Key Skills They Value")
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = ACCENT
        p = doc.add_paragraph(", ".join(a.get("key_skills_they_value")))
        doc.add_paragraph()

    add_sect("07", "Salary Expectation (INR)", a.get("salary_expectation_inr"))
    add_sect("08", "Growth Opportunities", a.get("growth_opportunities"))
    add_sect("09", "Risks & Challenges", a.get("risks"))

    if a.get("preparation_tips"):
        h = doc.add_paragraph()
        r = h.add_run("10  Preparation Tips")
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = ACCENT
        for i, tip in enumerate(a["preparation_tips"]):
            p = doc.add_paragraph(f"{i+1}. {tip}")
            p.paragraph_format.space_after = Pt(3)
        doc.add_paragraph()

    if a.get("profile_match"):
        h = doc.add_paragraph()
        r = h.add_run("WHY YOU MATCH THIS COMPANY")
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = GREEN
        p = doc.add_paragraph(str(a["profile_match"]))
        doc.add_paragraph()

    for text, size, color, bold in [
        ("RESUMATE AI", 9, ACCENT, True),
        ("Company Intelligence Engine", 8, MUTED, False),
        (f"Report generated on {now} at {time_str}", 8, MUTED, False),
    ]:
        fp = doc.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_after = Pt(2)
        r = fp.add_run(text)
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.font.bold = bold

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
